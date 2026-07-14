# -*- coding: utf-8 -*-
"""공모전 1차 서류심사 원고 생성 — 붙임2 편집기준(휴먼명조·여백·줄간격 160%) 준수 DOCX.
한글(HWP)에서 열어 '[서비스 아이디어 제안] 공모작명(신청자명).hwp'로 저장 후 제출.
개인정보 기재 금지 준수 — 이름·연락처 없음. 심사기준: 적합성40·창의성30·기대효과30.
사용: py -3.12 make_submission_manuscript.py
"""
import io, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
HERE = Path(__file__).parent
SCRATCH = Path(r"C:\Users\ain06\AppData\Local\Temp\claude\c--Users-ain06-OneDrive----2026---------------\0ddd6241-69e9-48f8-afe6-eeb71a4cf0cf\scratchpad")
OUT = HERE/"공모전자료"/"원고_LibAR_서비스아이디어제안_초안.docx"
DEMO_URL = "https://ainsof.dev/libar-demo/"
FONT = "휴먼명조"

# ── QR 코드 생성 ──
import qrcode
qr_path = SCRATCH/"demo_qr.png"
qrcode.make(DEMO_URL).save(str(qr_path))

doc = Document()

# 편집용지: A4, 여백 위25/아래20/좌30/우30, 머리말25/꼬리말20
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.top_margin, sec.bottom_margin = Mm(25), Mm(20)
sec.left_margin, sec.right_margin = Mm(30), Mm(30)
sec.header_distance, sec.footer_distance = Mm(25), Mm(20)

def set_font(run, size, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)

def para(text="", size=11, bold=False, align=None, indent=False, space_after=4,
         line=1.6, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None: p.alignment = align
    if indent: p.paragraph_format.first_line_indent = Pt(size*2)   # 첫 줄 들여쓰기 2자
    r = p.add_run(text)
    set_font(r, size, bold, color)
    return p

C, L = WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT

def heading(numeral, title):          # 대목차 Ⅰ. (휴먼명조 15 bold)
    para(f"{numeral}. {title}", size=15, bold=True, space_after=8)

def sub(no, title):                   # 중목차 1. (13)
    para(f"{no}. {title}", size=13, bold=True, space_after=6)

def subsub(ga, title):                # 소목차 가. (11)
    para(f"{ga}. {title}", size=11, bold=True, space_after=4)

def body(text, **kw):
    para(text, size=11, indent=True, **kw)

def bullet(text, size=11):
    p = para(f"○ {text}", size=size, space_after=3)
    p.paragraph_format.left_indent = Pt(12)
    return p

def caption(text):
    para(text, size=9, align=C, space_after=8)

def add_table(headers, rows, widths=None, size=10):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].alignment = C
        r = cell.paragraphs[0].add_run(h); set_font(r, size, bold=True)
        sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), "EFEFEF")
        cell._tc.get_or_add_tcPr().append(sh)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.paragraphs[0].alignment = C if j == 0 or len(str(v)) < 14 else L
            r = cell.paragraphs[0].add_run(str(v)); set_font(r, size)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows: row.cells[j].width = Mm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def add_image(path, width_mm, center=True):
    p = doc.add_paragraph()
    if center: p.alignment = C
    p.add_run().add_picture(str(path), width=Mm(width_mm))
    return p

def image_row(paths, width_mm):
    t = doc.add_table(rows=1, cols=len(paths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, pth in enumerate(paths):
        c = t.rows[0].cells[j]
        c.paragraphs[0].alignment = C
        c.paragraphs[0].add_run().add_picture(str(pth), width=Mm(width_mm))

def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# 꼬리말 쪽번호
footer_p = sec.footer.paragraphs[0]
footer_p.alignment = C
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
footer_p._p.append(fld)

# ═══════════ 표지 ═══════════
for _ in range(5): para()
para("2026 도서관 데이터 활용 공모전", size=15, align=C)
para("[서비스 아이디어 제안 부문]", size=13, align=C)
for _ in range(3): para()
para("LibAR", size=26, bold=True, align=C)
para("서가를 비추면 책이 제자리를 찾는다", size=18, bold=True, align=C)
para("— 도서관 데이터 기반 온디바이스 AI 서가 관리·탐색 서비스 —", size=13, align=C)
for _ in range(4): para()
add_image(qr_path, 30)
caption(f"실제 동작하는 공개 데모 (스마트폰 브라우저에서 바로 실행) · {DEMO_URL}")
page_break()

# ═══════════ 요약문 (1매 필수) ═══════════
para("요 약 문", size=15, bold=True, align=C, space_after=10)
add_table(["구분", "내용"], [
    ["기획의도 및\n내용 요약",
     "이용자가 찾는 책이 서가 제자리에 없으면 그 책은 '소장 중이지만 없는 책'이 된다. 오배열 점검은 사서가 "
     "청구기호를 한 권씩 눈으로 대조하는 수작업으로, 상시 수행이 어렵다. LibAR는 스마트폰 카메라로 서가를 "
     "비추기만 하면 AI가 책등 라벨(청구기호)을 읽어 장서 데이터와 대조하고, 잘못 꽂힌 책은 빨간색, 정상 책은 "
     "초록색으로 화면에 표시하는 웹 서비스다. 이용자에게는 찾는 책의 위치를 짚어 주는 도서찾기·길안내를 제공한다. "
     "설치 없이 브라우저에서 동작하며 전 과정이 기기 안에서 처리된다(공개 데모 운영 중)."],
    ["활용\n도서관 데이터",
     "① 도서관 정보나루(필수): 장서·대출 Open API — 실증관(영등포구립 대림도서관) 장서 목록 구축과 대출 상태 "
     "대조에 활용. 대출 중으로 기록된 책이 서가에서 발견되는 상태 불일치를 실증에서 7권 자동 검출. "
     "② 실증관 장서 데이터(공개 서지사항 1,860권): 청구기호 대조·오배열 판정의 기준 데이터."],
    ["활용 기술",
     "라벨 검출(YOLO 계열 경량 모델, 자가 학습 데이터 7,000여 박스), 청구기호 인식(PaddleOCR 한국어 모델을 "
     "저해상 라벨 실측 데이터로 파인튜닝), 장서 대조(퍼지 매칭+제목 이중 대조), 오배열 판정(최장증가부분수열 LIS), "
     "온디바이스 실행(ONNX Runtime Web, WebGPU/WASM). 모든 AI 모델은 부록에 명시."],
    ["추진 과정",
     "6주 실증: 실증관 4개 서가 구간(600~900번대) 사진·동영상 수집 → 인식 모델 파인튜닝(광각 라벨 판독률 6%→30%, 누적 고유 52권) →"
     "걷기 스캔 실증(동영상 30초로 서가 한 줄) → 판독률·확인율 공식 채점 → 브라우저 이식·공개 배포 → "
     "이용자 스캔이 학습 데이터가 되는 자가 개선 루프 가동."],
    ["결과물",
     "공개 웹 데모(설치 없이 즉시 체험), 걷기 스캔 판독률 95%·확인율 97%(151프레임 공식 채점), 실제 오배열 "
     "3건 검출(육안 확인), 대출-서가 상태 불일치 7권 검출, 전 과정 온디바이스(개인정보·사진 무전송)."],
    ["기대효과",
     "사서: 서가 한 줄 점검이 '30초 걷기'로 단축, 오배열 로그 축적. 이용자: '못 찾는 책' 감소, 서가 앞 실시간 "
     "길안내. 도서관: RFID 같은 고가 인프라 없이 스마트폰만으로 도입, 정보나루 연계로 전국 어느 도서관에나 "
     "장서 데이터만 연결하면 확장 가능. 오배열 로그×대출 빅데이터 매쉬업으로 서가 운영 분석의 새 데이터 창출."],
], widths=[28, 122], size=9.5)
page_break()

# ═══════════ 목차 ═══════════
para("목 차", size=15, bold=True, align=C, space_after=12)
for line in [
    "Ⅰ. 기획 배경 및 문제 정의",
    "Ⅱ. 서비스 제안: LibAR",
    "Ⅲ. 활용 데이터",
    "Ⅳ. 기술 구현 및 성능 검증",
    "Ⅴ. 추진 과정",
    "Ⅵ. 기대효과 및 발전 방향",
    "참고문헌",
    "부록. 공개 데모·산출물 및 사용 AI 기술 명세",
]:
    para(line, size=12, space_after=8)
page_break()

# ═══════════ Ⅰ. 기획 배경 ═══════════
heading("Ⅰ", "기획 배경 및 문제 정의")
sub("1", "서가 위 '잃어버린 책' — 오배열 문제")
body("도서관 자료는 청구기호 순서로 배가(排架)될 때에만 찾을 수 있다. 이용자든 사서든 책을 찾는 행위는 "
     "결국 '청구기호가 가리키는 자리'를 찾아가는 일이기 때문에, 한 권이 다른 자리에 꽂히는 순간 그 책은 "
     "목록에는 존재하지만 서가에서는 찾을 수 없는 책이 된다. 대출 후 반납 과정, 이용자가 열람 후 임의로 "
     "꽂아 두는 행동, 서가 정리 중의 착오 등으로 오배열은 도서관 운영에서 상시적으로 발생한다.")
body("본 팀이 실증 대상관(영등포구립 대림도서관)의 4개 서가 구간(600~900번대, 장서 1,860권 표본)을 "
     "스캔한 결과, 실제 오배열 3건이 검출되어 육안으로 확인되었다. 표본 규모를 감안하면 전체 서가에는 "
     "상시적으로 수십 건 규모의 오배열이 존재할 것으로 추정된다.")
sub("2", "기존 점검 방식의 한계")
bullet("육안 점검: 사서가 서가 앞에서 청구기호를 한 권씩 눈으로 대조 — 집중도가 필요한 반복 노동으로, "
       "다른 업무와 병행하기 어려워 점검 주기가 길어진다.")
bullet("RFID 기반 스마트 서가: 오배열을 자동 감지하지만 태그 부착·전용 서가·리더 등 권당·서가당 도입 "
       "비용이 커서 일부 도서관 외에는 확산되지 못했다.")
bullet("결과적으로 대다수 공공도서관에서 오배열 점검은 '가끔, 몰아서' 하는 작업이 되고, 그 사이 이용자는 "
       "소장 중인 책을 찾지 못한 채 돌아간다.")
sub("3", "착안점 — 스마트폰 카메라와 도서관 데이터의 결합")
body("모든 책등에는 이미 규격화된 청구기호 라벨이 붙어 있고, 도서관은 이미 완전한 장서 데이터를 보유하고 "
     "있으며, 사서와 이용자의 주머니에는 이미 고성능 카메라와 AI 연산이 가능한 스마트폰이 있다. LibAR는 "
     "이 세 가지 기존 자원을 소프트웨어로만 연결한다. 새 하드웨어 없이, 카메라로 서가를 비추면 AI가 라벨을 "
     "읽고 장서 데이터와 대조해 '이 자리에 있어야 할 책인가'를 즉시 판정하는 것이다.")

# ═══════════ Ⅱ. 서비스 제안 ═══════════
heading("Ⅱ", "서비스 제안: LibAR")
sub("1", "서비스 개요")
body("LibAR(Library + AR)는 스마트폰 브라우저에서 실행되는 서가 관리·탐색 서비스다. 설치가 필요 없고, "
     "촬영 영상이 서버로 전송되지 않으며(전 과정 온디바이스), 두 사용자 그룹에 각각의 화면을 제공한다.")
add_table(["사용자", "기능", "동작"], [
    ["사서", "서가 점검", "카메라로 서가를 훑으면 실시간으로 라벨 검출 → 일시정지하면 해당 화면의 청구기호를 "
     "판독·장서 대조 → 잘못 꽂힌 책 빨강, 확인된 책 초록으로 표시"],
    ["이용자", "도서찾기·길안내", "책 검색 → 층별 실배치도로 해당 서가까지 안내 → 서가 앞에서 스캔하면 "
     "찾는 책 위치를 화면에 짚어 줌"],
], widths=[20, 30, 100], size=10)
para("[그림 1] LibAR 앱 화면 — 홈, 서가 판독 결과(12권 확인), 도서찾기 길안내(층별 배치도)", size=9, align=C)
image_row([SCRATCH/"shot_home.png", SCRATCH/"shot_scan.png", SCRATCH/"shot_guide.png"], 44)
caption("출처: 본 팀 제작 공개 데모 화면 캡처")
sub("2", "이용 시나리오")
subsub("가", "사서 — 서가 점검")
body("반납 정리를 마친 사서가 스마트폰으로 서가 한 줄을 천천히 비추며 걷는다(약 30초). 화면 하단에 "
     "확인/잘못 꽂힘/다시 확인 권수가 실시간 집계되고, 잘못 꽂힌 책은 실물 위치에 빨간 박스로 표시되어 "
     "그 자리에서 바로잡는다. 판독 결과는 오배열 로그로 축적되어 어느 구간에서 오배열이 잦은지 통계를 만든다.")
subsub("나", "이용자 — 도서찾기")
body("이용자가 책 제목을 검색하면 소장 층·자료실과 서가 구간을 실배치도 위에 표시한다. 서가 앞에 도착해 "
     "'서가 스캔'을 누르고 카메라를 비추면, 찾는 책이 화면에 강조 표시된다. 오배열로 자리에 없는 경우에도 "
     "최근 스캔 로그에 위치가 남아 있으면 그 위치를 안내한다.")
sub("3", "차별성")
bullet("하드웨어 제로: RFID·스마트 서가와 달리 태그·전용 설비가 전혀 필요 없다. 스마트폰과 장서 데이터만으로 동작한다.")
bullet("온디바이스 처리: 영상·사진이 기기 밖으로 나가지 않아 이용자 프라이버시와 관내 보안 우려가 없고, 서버 비용도 없다.")
bullet("쓸수록 똑똑해지는 구조: 사용자가 동의하면 스캔한 라벨 조각(서가 전체 사진 아님)이 학습 데이터로 "
       "자동 수집되어 인식 모델이 현장 데이터로 계속 개선된다 — 도서관 현장이 곧 학습장이 되는 자가 개선 루프.")
bullet("제안이 아니라 실물: 본 원고의 모든 수치는 실증관에서 측정된 값이며, 공개 데모가 현재 운영 중이다"
       f"(부록, {DEMO_URL}).")

# ═══════════ Ⅲ. 활용 데이터 ═══════════
heading("Ⅲ", "활용 데이터")
sub("1", "도서관 데이터 (필수 데이터)")
add_table(["데이터", "출처", "활용 내용"], [
    ["장서·대출 Open API", "도서관 정보나루\n(data4library.kr)", "실증관 장서 목록 구축, 대출 상태 조회 — "
     "서가 스캔 결과와 대조해 '대출 중인데 서가에 있는 책'(상태 불일치) 자동 검출"],
    ["장서 서지 데이터", "실증관(대림도서관)\n공개 서지사항", "청구기호·서명·자료실 1,860권 — 인식 결과 대조, "
     "오배열 판정, 도서찾기 검색·길안내의 기준 데이터"],
    ["전국 대출 빅데이터", "도서관 정보나루", "발전 단계(Ⅵ장): 오배열 로그와 매쉬업하여 '많이 찾는 책일수록 "
     "오배열이 잦은가' 등 서가 운영 분석"],
], widths=[32, 36, 82], size=10)
sub("2", "데이터가 실제로 일한 사례 — 대출·서가 상태 불일치 검출")
body("정보나루 대출 데이터는 하루 단위(D-1)로 갱신된다. 실증 중 대출 중으로 기록된 책이 서가에서 스캔으로 "
     "발견되는 사례가 누적 7권 확인되었는데, 이는 반납 처리와 배가 사이의 시차·착오를 데이터로 잡아낸 것이다. "
     "즉 LibAR는 서가(실물)와 데이터(기록)의 양방향 검증 장치로 동작하며, 이것이 본 서비스가 도서관 데이터를 "
     "'조회'가 아니라 '대조'에 활용하는 핵심 방식이다.")
sub("3", "서비스가 만드는 새 데이터")
body("판독 로그(어느 서가에서 몇 권이 확인·오배열되었는지)와 옵트인 라벨 조각(정답 청구기호가 붙은 저해상 "
     "라벨 이미지)이 축적된다. 전자는 서가 운영 통계의 원천이 되고, 후자는 인식 모델 재학습 데이터가 된다. "
     "실증에서 실전 라벨 조각 702줄로 학습셋을 구축했고, 현장 테스트의 조각이 공유 저장소로 자동 수집되고 있다.")

# ═══════════ Ⅳ. 기술 구현 ═══════════
heading("Ⅳ", "기술 구현 및 성능 검증")
sub("1", "인식 파이프라인 — 4단계")
add_table(["단계", "역할", "구현"], [
    ["① 검출", "서가 영상에서 청구기호 라벨 위치 탐지", "YOLO 계열 경량 모델 — 파이프라인이 검증한 라벨 "
     "박스 7,000여 개를 자동 라벨로 학습(수작업 라벨링 0회)"],
    ["② 인식", "라벨의 청구기호 문자 판독", "PaddleOCR 한국어 인식 모델을 실측 저해상 라벨로 파인튜닝 — "
     "광각 판독률 6%→30%(동일 조건 A/B, 5배), 개선 누적 시 고유 52권"],
    ["③ 대조", "판독 결과를 장서 데이터와 매칭", "분류번호·저자기호 퍼지 매칭 + 판독 실패 시 책등 제목 "
     "이중 대조로 복구"],
    ["④ 판정", "행 안의 순서로 오배열 식별", "최장증가부분수열(LIS) — 최소 오배열 집합만 지목해 이웃 오탐 방지"],
], widths=[16, 48, 86], size=10)
para("[그림 2] 근접 스캔 판정 결과 — 확인(초록)·오배열(빨강)·제목복구(파랑)", size=9, align=C)
add_image(HERE/"리포트이미지"/"근접_01_ar.jpg", 130)
caption("출처: 실증관 서가 촬영·본 팀 파이프라인 산출")
sub("2", "성능 — 3지표 공식 채점")
body("서비스 품질은 '몇 권 인식'이 아니라 세 지표로 측정했다. 걷기 스캔 동영상 151프레임을 근접 촬영으로 "
     "확정한 정답지 87권과 교차 채점한 결과다.")
add_table(["지표", "정의", "실측값"], [
    ["판독률", "서가에 실재하는 책 중 정확히 읽어낸 비율(재현율)", "95%"],
    ["확인율", "읽었다고 표시한 것 중 실제로 맞은 비율(정밀도)", "97%"],
    ["오배열 검증", "검출된 오배열의 육안 확인", "실제 오배열 3건 확인, 오탐 0건 유지"],
], widths=[26, 90, 34], size=10)
body("검출을 놓친 책도 이웃 프레임 투표로 회수된다. 같은 서가를 30초 동영상으로 훑으면 프레임당 6~7권이 "
     "투표 합산으로 80권 이상으로 증폭되는 것을 실측했다 — 걷기만 해도 점검이 완성되는 근거다.")
para("[그림 3] 걷기 스캔 프레임 — 동영상에서 실시간 검출·판독", size=9, align=C)
add_image(HERE/"libar-sample"/"hybrid_v3_results"/"yolo_동영상1_f00060_ar.jpg", 130)
caption("출처: 실증관 걷기 스캔 동영상 프레임·본 팀 파이프라인 산출")
sub("3", "온디바이스 실행 — 서버 없는 AI")
body("검출·인식·대조·판정 전 과정을 브라우저 안에서 실행하도록 모델을 경량화(ONNX, 반정밀도)해 이식했다. "
     "GPU 가속(WebGPU)이 되는 기기에서는 서가 사진 한 장을 약 20초에, 미지원 기기도 자동 전환(WASM)으로 "
     "약 30초에 판독한다. 사진·영상이 기기를 떠나지 않으므로 프라이버시 문제가 원천 차단되고, 도서관은 "
     "서버 운영 비용 없이 웹 주소 하나로 서비스를 도입한다.")
sub("4", "자가 개선 루프 — 현장이 학습장이 된다")
body("인식 모델의 성능은 실제 현장 라벨 데이터에 좌우된다. LibAR는 사용자가 1회 동의하면 스캔할 때마다 "
     "라벨 조각(정답 청구기호 포함, 서가 전체 사진 제외)을 학습 수집소로 자동 전송한다. 수집된 조각은 "
     "자동 변환기를 거쳐 학습쌍이 되고, 재학습된 모델이 다시 앱으로 배포된다. 실증 기간에 이 루프로 "
     "파인튜닝 4세대를 거치며 광각 판독률을 6%에서 30%로(동일 조건), 개선 누적으로 고유 52권까지 끌어올렸다.")
para("[그림 4] 자가 학습 데이터 예 — 실측 라벨 조각과 정답 청구기호 쌍", size=9, align=C)
add_image(HERE/"libar-sample"/"real_rec_data_v3"/"pair_montage_3rd.jpg", 140)
caption("출처: 실증 수집 데이터로 본 팀 생성")

# ═══════════ Ⅴ. 추진 과정 ═══════════
heading("Ⅴ", "추진 과정")
add_table(["단계", "기간", "내용", "결과"], [
    ["M1 측정", "1주차", "실증관 서가 촬영(근접·광각), 기준 지표 수립", "근접 매칭 90%, 광각 한계 확인(라벨 40px)"],
    ["M2 일반화", "2주차", "새 서가 구간(색상 다른 라벨) 추가 비용 검증", "구간 추가 수정 0~2줄 — 일반화 성립"],
    ["M3 하이브리드", "3~4주차", "AI 검출기 자가 학습(자동 라벨) + 파이프라인 통합", "전 지형에서 규칙 기반 대비 우세"],
    ["M4 걷기 스캔", "5주차", "동영상 스캔·프레임 투표, 판독률·확인율 공식 채점", "판독률 95%·확인율 97%"],
    ["M5 온디바이스", "6주차", "모델 경량화·브라우저 이식, 공개 배포", "공개 데모 운영, 전 과정 기기 내 처리"],
], widths=[24, 18, 66, 42], size=9.5)
body("전 과정에서 '기각 실험'도 기록으로 남겼다. 초해상도 전처리(효과 0), 개별 크롭 인식(검출 붕괴), "
     "과도한 박스 타이트화(정확도 후퇴) 등 실패한 접근을 데이터로 판정하고 되돌렸으며, 이는 본 제안의 "
     "수치가 선별된 성공 사례가 아니라 재현 가능한 측정값임을 뒷받침한다.")

# ═══════════ Ⅵ. 기대효과 ═══════════
heading("Ⅵ", "기대효과 및 발전 방향")
sub("1", "기대효과")
subsub("가", "사서 — 점검 노동의 구조 전환")
body("서가 한 줄 점검이 '청구기호 육안 대조 수십 분'에서 '30초 걷기 + 빨간 박스 확인'으로 바뀐다. "
     "점검이 가벼워지면 주기가 짧아지고, 오배열이 빨리 잡히면 이용자가 책을 못 찾는 시간도 줄어드는 "
     "선순환이 생긴다. 축적되는 판독 로그는 구간별 오배열 통계라는, 지금까지 존재하지 않던 운영 데이터를 만든다.")
subsub("나", "이용자 — 정보 접근성 개선")
body("'목록에는 있는데 서가에 없는 책'은 이용자 입장에서 사실상의 결본이다. 오배열이 줄고, 서가 앞 "
     "실시간 안내가 더해지면 도서관 이용 문턱이 낮아진다. 특히 청구기호 체계에 익숙하지 않은 이용자"
     "(어린이, 고령층, 도서관 초심자)에게 '카메라로 비추면 짚어 주는' 안내는 정보 격차를 직접 줄인다.")
subsub("다", "도서관 — 저비용 확산")
body("필요한 것은 장서 데이터(모든 도서관이 보유)와 스마트폰(모든 사서·이용자가 보유)뿐이다. 정보나루 "
     "Open API로 장서 데이터를 연결하면 타 도서관 확장 시 개발 비용이 거의 들지 않으며, 실증에서 새 서가 "
     "구간 추가에 수정 0~2줄임을 확인했다.")
sub("2", "발전 방향")
bullet("전국 확장: 정보나루 API 기반 도서관 선택형 서비스로 확대 — 도서관별 장서 데이터 연결만으로 즉시 적용.")
bullet("오배열 로그 × 대출 빅데이터 매쉬업: '대출이 잦은 주제일수록 오배열도 잦은가', '오배열 다발 구간과 "
       "서가 배치의 관계' 등 정보나루 전국 대출 데이터와 결합한 서가 운영 분석 — 도서관 데이터 생태계에 "
       "'서가 실물 상태'라는 새 데이터 축을 추가한다.")
bullet("장서 점검(재고 조사) 지원: 걷기 스캔 로그를 누적하면 연 단위 장서 점검의 사전 스크리닝으로 활용 가능.")
bullet("접근성 특화: 시각 약자를 위한 음성 안내('세 번째 칸, 왼쪽에서 다섯 번째'), 다국어 이용자를 위한 "
       "서명 표시 등 온디바이스 특성을 살린 확장.")
page_break()

# ═══════════ 참고문헌 ═══════════
heading("참고문헌", "")
doc.paragraphs[-1].runs[0].text = "참고문헌"
for ref in [
    "국립중앙도서관. (2026). 2026 도서관 데이터 활용 공모전 안내문. 국립중앙도서관.",
    "도서관 정보나루. (2026). 도서관 빅데이터 Open API. https://www.data4library.kr",
    "영등포구립도서관. (2026). 대림도서관 자료실 안내 및 소장자료 검색. https://www.ydplib.or.kr",
    "Jocher, G., et al. (2024). Ultralytics YOLO (Version 11/26) [Computer software]. https://github.com/ultralytics/ultralytics",
    "PaddlePaddle. (2025). PaddleOCR: Awesome multilingual OCR toolkits (PP-OCRv5) [Computer software]. https://github.com/PaddlePaddle/PaddleOCR",
    "Microsoft. (2025). ONNX Runtime Web [Computer software]. https://onnxruntime.ai",
]:
    p = para(ref, size=10, space_after=4)
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.first_line_indent = Pt(-20)

# ═══════════ 부록 ═══════════
para()
heading("부록", "공개 데모·산출물 및 사용 AI 기술 명세")
sub("1", "공개 데모")
body(f"주소: {DEMO_URL} (스마트폰·PC 브라우저에서 즉시 실행, 설치 불필요). '샘플 서가' 버튼으로 실증관 "
     "서가 사진 판독 전 과정을, '서가 점검하기'로 실시간 카메라 검출을 체험할 수 있다. 촬영 영상은 "
     "전송되지 않는다.")
add_image(qr_path, 26)
sub("2", "사용 AI 기술 명세 (요강의 AI 기술 명시 의무 준수)")
add_table(["구분", "기술", "용도·산출물"], [
    ["검출 모델", "Ultralytics YOLO 경량 모델(nano)을 자체 데이터 7,000여 박스로 전이학습",
     "서가 영상의 청구기호 라벨 위치 검출(ONNX 변환 탑재)"],
    ["인식 모델", "PaddleOCR korean PP-OCRv5 인식 모델을 실측 라벨 702줄+합성 데이터로 파인튜닝",
     "청구기호 문자 판독(ONNX 반정밀도 변환 탑재)"],
    ["텍스트 줄 검출", "PaddleOCR PP-OCRv5 mobile 검출 모델(사전학습 그대로)", "라벨 내 텍스트 줄 분리"],
    ["실행 환경", "ONNX Runtime Web (WebGPU/WASM)", "브라우저 온디바이스 추론"],
    ["개발 보조", "AI 코딩 어시스턴트(Anthropic Claude)", "코드 작성·실험 자동화 보조 — 모든 실험 설계·"
     "판정·실측은 본 팀이 수행·검증"],
], widths=[26, 62, 62], size=9.5)
sub("3", "산출물 목록")
bullet("공개 웹 데모 및 소스(공개 저장소), 검출·인식 모델(ONNX), 판독률·확인율 채점 스크립트")
bullet("실증 데이터: 서가 구간 4곳 1,860권 장서 대조 데이터, 걷기 스캔 151프레임 채점 결과, 학습 라벨 7,000여 박스")
bullet("2차 발표심사 시 '분석에 활용한 데이터 및 모델 1식' 제출 가능(모델 파일·학습 데이터·채점 코드)")

doc.save(str(OUT))
print(f"저장: {OUT}")
print("다음 단계: 한글(HWP)에서 열어 서식 확인 → '[서비스 아이디어 제안] LibAR(신청자명).hwp'로 저장")
