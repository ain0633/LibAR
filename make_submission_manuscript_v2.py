# -*- coding: utf-8 -*-
"""공모전 1차 서류심사 원고 [초안2] — 비전공 심사위원용 쉬운 버전 (~10페이지).
초안1(상세판)과 별도 문서. 붙임2 편집기준(휴먼명조·여백·줄간격 160%) 동일 준수.
원칙: 합니다체 / 전문용어는 부록으로 / 핵심 숫자 6개만 / 한 절 = 한 메시지.
사용: py -3.12 make_submission_manuscript_v2.py
"""
import io, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
HERE = Path(__file__).parent
SCRATCH = Path(r"C:\Users\ain06\AppData\Local\Temp\claude\c--Users-ain06-OneDrive----2026---------------\0ddd6241-69e9-48f8-afe6-eeb71a4cf0cf\scratchpad")
OUT = HERE/"공모전자료"/"원고_LibAR_서비스아이디어제안_초안2.docx"
DEMO_URL = "https://ainsof.dev/libar-demo/"
FONT = "휴먼명조"

import qrcode
qr_path = SCRATCH/"demo_qr.png"
qrcode.make(DEMO_URL).save(str(qr_path))

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.top_margin, sec.bottom_margin = Mm(25), Mm(20)
sec.left_margin, sec.right_margin = Mm(30), Mm(30)
sec.header_distance, sec.footer_distance = Mm(25), Mm(20)

def set_font(run, size, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)

def para(text="", size=11, bold=False, align=None, indent=False, space_after=4, line=1.6):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None: p.alignment = align
    if indent: p.paragraph_format.first_line_indent = Pt(size*2)
    r = p.add_run(text); set_font(r, size, bold)
    return p

C = WD_ALIGN_PARAGRAPH.CENTER
L = WD_ALIGN_PARAGRAPH.LEFT

def heading(numeral, title): para(f"{numeral}. {title}", size=15, bold=True, space_after=8)
def sub(no, title): para(f"{no}. {title}", size=13, bold=True, space_after=6)
def body(text, **kw): para(text, size=11, indent=True, **kw)
def bullet(text, size=11):
    p = para(f"○ {text}", size=size, space_after=3)
    p.paragraph_format.left_indent = Pt(12)
def caption(text): para(text, size=9, align=C, space_after=8)

def add_table(headers, rows, widths=None, size=10):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].alignment = C
        r = cell.paragraphs[0].add_run(h); set_font(r, size, bold=True)
        sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), "EFEFEF")
        cell._tc.get_or_add_tcPr().append(sh)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.paragraphs[0].alignment = C if j == 0 or len(str(v)) < 14 else L
            r = cell.paragraphs[0].add_run(str(v)); set_font(r, size)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows: row.cells[j].width = Mm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_image(path, width_mm, center=True):
    p = doc.add_paragraph()
    if center: p.alignment = C
    p.add_run().add_picture(str(path), width=Mm(width_mm))

def image_row(paths, width_mm):
    t = doc.add_table(rows=1, cols=len(paths))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, pth in enumerate(paths):
        c = t.rows[0].cells[j]
        c.paragraphs[0].alignment = C
        c.paragraphs[0].add_run().add_picture(str(pth), width=Mm(width_mm))

def banner(text, fill="D9EAD3"):                       # 기능 배너 (수상작 스타일 — 연녹색 띠)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Mm(150)
    cell.paragraphs[0].alignment = C
    r = cell.paragraphs[0].add_run(text); set_font(r, 11.5, bold=True)
    sh = OxmlElement("w:shd"); sh.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(sh)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def captioned_row(items, width_mm):                    # [(이미지, 캡션)] — 화면 흐름 + 화면별 설명
    t = doc.add_table(rows=2, cols=len(items))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (pth, cap) in enumerate(items):
        c = t.rows[0].cells[j]
        c.paragraphs[0].alignment = C
        c.paragraphs[0].add_run().add_picture(str(pth), width=Mm(width_mm))
        c2 = t.rows[1].cells[j]
        c2.paragraphs[0].alignment = C
        r = c2.paragraphs[0].add_run(cap); set_font(r, 8.5, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(): doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

footer_p = sec.footer.paragraphs[0]
footer_p.alignment = C
fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
footer_p._p.append(fld)

# ═══════════ 표지 ═══════════
for _ in range(5): para()
para("2026 도서관 데이터 활용 공모전", size=15, align=C)
para("[서비스 아이디어 제안 부문]", size=13, align=C)
for _ in range(3): para()
para("LibAR", size=26, bold=True, align=C)
para("도서관에서 가장 찾기 힘든 책은,", size=18, bold=True, align=C)
para("잘못 꽂힌 책입니다", size=18, bold=True, align=C)
para("— 카메라를 비추는 순간, AI가 20,944권과 대조해 책의 제자리를 찾아 줍니다 —", size=13, align=C)
for _ in range(4): para()
add_image(qr_path, 30)
caption(f"실제로 동작하는 공개 데모 — 스마트폰으로 QR을 찍으면 바로 체험할 수 있습니다 · {DEMO_URL}")
page_break()

# ═══════════ 요약문 (1매 필수) ═══════════
para("요 약 문", size=15, bold=True, align=C, space_after=10)
add_table(["구분", "내용"], [
    ["기획의도 및\n내용 요약",
     "도서관 책은 청구기호 순서대로 꽂혀 있을 때만 찾을 수 있습니다. 그런데 이용자에게 청구기호는 낯선 "
     "암호이고, 한 권이 잘못 꽂히면 그 책은 '소장 중인데 찾을 수 없는 책'이 됩니다. 사서의 서가 점검은 "
     "한 권씩 눈으로 대조하는 수작업이라 자주 하기 어렵습니다. LibAR는 스마트폰 카메라로 서가를 비추면 "
     "AI가 책등의 청구기호를 읽어 도서관 장서 목록과 맞춰 보는 서비스입니다. 이용자에게는 찾는 책 한 권을 "
     "화면에 짚어 주고 서가 위에서 사서추천·인기대출 도서까지 발견하게 하며, 사서에게는 잘못 꽂힌 책을 "
     "빨간색으로 표시해 줍니다. 설치 없이 웹 주소만 열면 되고, 촬영 화면은 스마트폰 밖으로 전송되지 "
     "않습니다. 실증 도서관에서 검증을 마쳤고 공개 데모를 운영 중입니다."],
    ["활용\n도서관 데이터",
     "① 실증관 장서 데이터(영등포구립 대림도서관, 전 장서 20,944권 — 청구기호·서명·자료실·대출 상태): "
     "책 대조·오배열 판정·도서찾기의 기준. 대출 중으로 기록된 책이 서가에서 발견되는 불일치 7권을 자동으로 "
     "찾아냈습니다. ② 도서관 정보나루 Open API(인기대출도서): 최근 1년 전국 대출 빅데이터를 분류별로 받아 "
     "실증관 소장 357권과 연결 — 서가를 비추면 인기대출 도서가 🔥핀으로 표시되는 탐색 기능으로 구현했습니다. "
     "③ 국립중앙도서관 사서추천도서 Open API: 2016년 이후 추천도서를 실증관 소장 222권과 연결 — 서가 위에서 "
     "⭐핀을 누르면 추천사가 보입니다. ④ 장서 데이터 분석: 책 구성과 대출 상태로 '오배열이 잦을 구간'을 "
     "예측하는 위험도 분석 — 실제 문제가 발견된 9개 구간 중 7개를 예측이 맞혔습니다."],
    ["활용 기술",
     "책등 라벨의 위치를 찾는 AI(딥러닝 검출 모델)와 청구기호 글자를 읽는 AI(문자인식 모델)를 실증관 "
     "서가 사진으로 직접 훈련·개선했고, 스마트폰 브라우저 안에서 바로 실행되도록 압축(경량화)했습니다. "
     "읽은 결과를 장서 데이터와 대조하고 책 순서를 계산해 잘못 꽂힌 책을 가려냅니다. 서버가 없어 운영 "
     "비용이 0원이고, 개인정보 걱정이 원천적으로 없습니다. 사용한 AI 기술은 부록에 전부 명시했습니다."],
    ["추진 과정",
     "7주간 실증: 실증관 서가를 촬영해 데이터를 만들고 → AI를 훈련하고 → 서가를 걸으며 스캔하는 방식을 "
     "검증하고 → 판독 결과를 해석하는 논리를 다듬어 정확도를 끌어올리고 → 스마트폰에서 돌아가게 만들어 "
     "공개 배포한 뒤 → 현직 사서와 함께 현장에서 다듬었습니다. 효과가 없는 방법은 실험으로 확인해 "
     "버리고, 효과가 검증된 것만 남겼습니다."],
    ["결과물",
     "공개 웹 데모(설치 없이 즉시 체험). 서가를 걸으며 스캔하면 책의 99%를 정확히 읽어냅니다(정답지 "
     "87권 공식 채점). 사진 한 장이면 16권을 약 8초에 확인합니다. 실증에서 실제 오배열 4건, 전산 기록과 "
     "서가 실물이 다른 사례 6건, 대출 중인데 서가에 있는 책 7권을 찾아냈습니다."],
    ["기대효과",
     "이용자: 청구기호를 몰라도 카메라만 비추면 책을 찾습니다. 사서: 서가 한 줄 점검이 30초 걷기로 "
     "줄어듭니다. 도서관: 값비싼 장비(RFID 등) 없이 웹 주소 하나로 도입하고, 장서 데이터만 연결하면 "
     "전국 어느 도서관에서나 쓸 수 있습니다."],
], widths=[28, 122], size=9.5)
page_break()

# ═══════════ 목차 ═══════════
para("목 차", size=15, bold=True, align=C, space_after=12)
for line in [
    "Ⅰ. 왜 만들었나 — 책을 못 찾는 도서관",
    "Ⅱ. 무엇을 만들었나 — LibAR 서비스 소개",
    "Ⅲ. 도서관 데이터를 어떻게 활용했나",
    "Ⅳ. 어떻게 만들었나 — 기술과 성능",
    "Ⅴ. 추진 경과",
    "Ⅵ. 기대효과 및 발전 방향",
    "참고문헌",
    "부록. 공개 데모 및 사용 AI 기술 명세",
]:
    para(line, size=12, space_after=8)
page_break()

# ═══════════ Ⅰ ═══════════
heading("Ⅰ", "왜 만들었나 — 책을 못 찾는 도서관")
body("도서관에서 가장 찾기 힘든 책은 절판된 책도, 대출 중인 책도 아닙니다 — 잘못 꽂힌 책입니다. 목록에는 "
     "있는데 서가에서는 아무도 찾을 수 없기 때문입니다. 그리고 그 앞에는, 청구기호라는 암호를 풀지 못해 "
     "제자리에 있는 책조차 못 찾는 이용자가 있습니다.")
sub("1", "이용자: 청구기호는 어렵습니다")
body("검색으로 '813.6-김54ㅎ' 같은 청구기호를 얻어도, 그 암호를 풀어 서가 위치로 옮기는 일은 이용자 "
     "몫입니다. 실증관 종합자료실만 해도 한국소설(813) 구간 한 곳에 1,839권이 꽂혀 있습니다 — 청구기호를 "
     "못 읽는 이용자가 책 제목만 보고 훑을 수 있는 규모가 아닙니다. 어린이, 어르신, 도서관이 처음인 "
     "분들일수록 여기서 막히고, 찾다 포기한 경험은 다음 방문을 막습니다.")
sub("2", "사서: 잘못 꽂힌 책을 잡아낼 도구가 없습니다")
body("반납 정리, 이용자가 아무 데나 꽂아 둔 책, 정리 중의 실수로 오배열은 늘 생깁니다. 잘못 꽂힌 책은 "
     "목록에는 있지만 서가에서는 찾을 수 없는, 사실상 사라진 책입니다. 저희가 실증관 전 장서 20,944권을 "
     "기준 삼아 일부 구간만 표본 점검했는데도 실제 오배열 4건과, 전산 기록과 서가 실물이 어긋난 사례 "
     "6건이 나왔습니다. 일부에서 이만큼이면 전체에는 더 있다는 뜻인데, 점검은 여전히 눈으로 한 권씩 "
     "대조하는 수작업입니다. RFID 같은 자동화 장비는 비용 때문에 대다수 도서관이 도입하지 못했습니다.")
sub("3", "착안: 필요한 것은 이미 다 있습니다")
body("모든 책등에는 청구기호 라벨이 붙어 있고, 도서관에는 완전한 장서 데이터가 있고, 모두의 주머니에는 "
     "고성능 카메라가 달린 스마트폰이 있습니다. 이 셋을 소프트웨어로 연결하면 — 카메라로 서가를 비추는 "
     "것만으로 '이 책이 제자리에 있는가'를 즉시 알 수 있습니다. 같은 판정을 이용자에게 쓰면 도서찾기가 "
     "되고, 사서에게 쓰면 서가 점검이 됩니다. 새로 사야 할 장비는 없습니다.")

# ═══════════ Ⅱ ═══════════
heading("Ⅱ", "무엇을 만들었나 — 아이디어 실행 결과")
body("LibAR(Library + AR)는 스마트폰 브라우저에서 열리는 웹 서비스입니다. 설치가 필요 없고, 촬영 화면이 "
     "폰 밖으로 나가지 않으며, 이용자와 사서에게 각각의 화면을 제공합니다. 아래 화면은 모두 실제로 "
     "동작하는 공개 데모의 실행 화면입니다.")
banner("이용자 기능 — 도서찾기·길안내: 청구기호를 몰라도 책 앞까지")
captioned_row([
    (SCRATCH/"v2_find.png", "① 도서찾기 — 검색·추천도서·인기대출 선택지"),
    (SCRATCH/"v2_find_pop.png", "② 인기대출 목록 — 전국 대출 빅데이터 중 소장분"),
    (SCRATCH/"v2_guide.png", "③ 길안내 — 층별 지도로 서가까지"),
], 43)
captioned_row([
    (HERE/"도서찾기_핀_예시2.png", "④ 서가 앞에서 카메라를 비추면 — 찾는 책 한 권만 확률과 함께 📍 (실증관 실측 화면)"),
], 52)
bullet("화면에 표시되는 확률(98% 등)은 장식이 아니라 실측값입니다 — 자리 추정 정확도는 검증 273건 전부 "
       "적중(100%)이었습니다.")
bullet("찾는 책이 화면 밖이면 \"오른쪽으로 약 12권 옆이에요\"처럼 배가 순서 데이터로 계산한 이동 안내를 줍니다.")
captioned_row([
    (SCRATCH/"v2_discover_pop.png", "⑤ 서가를 비추면 도서관 데이터가 보입니다 — ⭐사서추천·🔥인기대출 핀을 누르면 추천사·대출 정보 (실증관 서가 실측 화면)"),
], 52)
bullet("탐색 기능: 정보나루 인기대출도서(소장 357권 연결)와 국립중앙도서관 사서추천도서(소장 222권 연결)를 "
       "서가 실물 위에 겹쳐 보여 줍니다 — 도서관 데이터가 이용자의 눈앞에서 책이 되는 순간입니다.")
banner("사서 기능 — 서가 점검: 30초 걷기, 사진 한 장이면 끝")
captioned_row([
    (SCRATCH/"v2_scan_ar.png", "① 스캔 — 확인은 초록, 잘못 꽂힘은 빨강"),
    (SCRATCH/"v2_reason.png", "② 근거 — 지금 자리 ↔ 올바른 자리, 재배열 완료 기록"),
    (SCRATCH/"v2_dash.png", "③ 점검 현황 — 기록이 자동으로 업무 데이터로"),
], 43)
bullet("사진 점검: 사진 한 장에 16권을 약 8초에 확인하고, 점검 사진은 리포트로 저장·전송되어 업무 기록이 "
       "됩니다. 가까이 찍은 단은 '확인', 멀어서 읽기 어려운 라벨은 '다시 확인'으로 구분해 알려 줍니다.")
bullet("잘못 꽂힌 책은 그 자리에서 바로잡고 '재배열 완료'를 누르면 조치까지 기록됩니다 — 점검이 곧 통계가 "
       "됩니다.")
sub("1", "다른 방법과 무엇이 다른가")
bullet("데이터가 실물 위에: 전국 대출 빅데이터와 국립중앙도서관 추천이 '목록 화면'이 아니라 눈앞의 서가 위 "
       "핀으로 보입니다 — 도서관 데이터를 이용자의 시선 높이로 가져온 새로운 인터페이스입니다.")
bullet("장비 0원: RFID·스마트 서가와 달리 태그도 전용 설비도 필요 없습니다. 스마트폰과 장서 데이터면 됩니다.")
bullet("개인정보 원천 차단: 모든 AI 처리가 폰 안에서 끝나 촬영 화면이 어디로도 전송되지 않습니다. 서버가 "
       "없으니 운영 비용도 0원입니다.")
bullet("제안이 아니라 실물: 이 원고의 모든 수치는 실증 도서관에서 측정한 값이고, 지금 공개 데모가 "
       f"운영 중입니다(부록 QR, {DEMO_URL}).")

# ═══════════ Ⅲ ═══════════
heading("Ⅲ", "도서관 데이터를 어떻게 활용했나")
sub("1", "활용 데이터")
add_table(["데이터", "출처", "쓰임새"], [
    ["장서 데이터\n(전 장서 20,944권)", "실증관\n(대림도서관)", "청구기호·서명·자료실·대출 상태 — 카메라가 읽은 "
     "책을 맞춰 보는 기준이자, 오배열 판정과 도서찾기의 뿌리"],
    ["인기대출도서\n(전국 대출 빅데이터)", "도서관 정보나루\nOpen API", "분류별 최근 1년 인기대출을 ISBN으로 실증관 "
     "소장과 연결(357권) — 서가 화면 위 🔥핀 탐색 기능. 발전 단계에는 오배열 기록과 결합한 점검 우선순위 "
     "추천으로 확장(Ⅵ장)"],
    ["사서추천도서", "국립중앙도서관\nOpen API", "2016년 이후 추천도서·추천사를 실증관 소장과 연결(222권) — "
     "서가 화면 위 ⭐핀·추천사 팝업"],
], widths=[36, 32, 82], size=10)
sub("2", "데이터가 실제로 일한 사례 두 가지")
body("첫째, 대조가 불일치를 잡았습니다. 서가를 스캔하다가 '전산에는 대출 중인데 서가에 꽂혀 있는 책' "
     "7권을 자동으로 찾아냈습니다. 반납 처리와 실제 배가 사이의 시차·착오가 데이터로 드러난 것입니다. "
     "전산 목록에 없는 실물, 라벨과 전산이 다른 책 등 6건도 같은 방식으로 발견됐습니다. 서가(실물)와 "
     "데이터(기록)를 양방향으로 검증하는 장치 — 이것이 본 서비스가 도서관 데이터를 쓰는 핵심 방식입니다.")
body("둘째, 데이터로 오배열을 예측해 봤습니다. '비슷한 책이 빽빽한 구간일수록, 손을 많이 타는 구간일수록 "
     "잘못 꽂히기 쉽다'는 가설로 177개 구간의 위험도 점수를 만들었더니 — 실증 중 실제 문제가 발견된 "
     "9개 구간 중 7개가 위험도 상위 절반에 들었습니다. 오배열 기록이 쌓이면 이 분석은 '점검 우선순위 "
     "추천'이 됩니다. 서가 점검을 감이 아니라 데이터로 계획하게 만드는 것입니다.")
sub("3", "서비스가 만드는 새 데이터")
body("점검할 때마다 판독 기록(어느 서가에서 몇 권 확인·오배열)과 라벨 조각 데이터가 쌓입니다. 실증 "
     "7주간 모인 조각 5,200여 개는 사람이 정답을 검증해, 모든 개선이 진짜 개선인지 판정하는 시험지"
     "(현장 평가셋)가 됐습니다. 점검 리포트는 사서의 업무 기록으로 저장됩니다.")

# ═══════════ Ⅳ ═══════════
heading("Ⅳ", "어떻게 만들었나 — 기술과 성능")
sub("1", "작동 원리 — 네 단계")
add_table(["단계", "하는 일"], [
    ["① 찾기", "서가 화면에서 청구기호 라벨의 위치를 AI가 찾습니다 (실증관 사진 7,000여 개 라벨로 직접 훈련)"],
    ["② 읽기", "라벨의 청구기호 글자를 AI가 읽습니다 (한국어 문자인식 모델을 실측 라벨로 추가 훈련)"],
    ["③ 맞추기", "읽은 결과를 장서 20,944권과 대조합니다. 글자가 일부 틀려도 책등 제목까지 함께 봐서 복구합니다"],
    ["④ 판정", "책들의 순서를 계산해 '순서가 어긋난 책'만 정확히 지목합니다 — 이웃 책까지 잘못 의심하지 않도록"],
], widths=[20, 130], size=10)
para("[그림 3] 판정 결과 — 확인(초록) · 오배열(빨강) · 제목으로 복구(파랑)", size=9, align=C)
add_image(HERE/"리포트이미지"/"근접_01_ar.jpg", 125)
caption("출처: 실증관 서가 촬영, 본 팀 산출")
sub("2", "성능 — 공식 채점 결과")
add_table(["항목", "결과"], [
    ["판독 정확도 (서가를 걸으며 스캔)", "99% — 정답지 87권 중 86권 (공식 채점)"],
    ["표시의 신뢰도 (확인이라고 표시한 것 중 실제 정답)", "97%"],
    ["사진 한 장 점검", "16권 확인 / 약 8초 (스마트폰 단독)"],
    ["못 읽은 책의 자리 추정 (도서찾기)", "정확도 100% — 검증 273건 전부 적중"],
], widths=[95, 55], size=10)
body("한 가지를 강조하고 싶습니다. 정확도를 89%에서 99%로 올린 마지막 도약은 AI를 더 학습시켜 얻은 것이 "
     "아닙니다. AI는 이미 답을 읽고 있었는데 그 결과를 해석하는 규칙이 정답을 버리고 있었고, 현장 실패 "
     "사례를 분석해 그 규칙을 고쳐 얻은 것입니다. 반대로 '더 학습시키면 좋아지겠지' 싶은 방법 5가지는 "
     "실험으로 전부 효과 없음을 확인하고 버렸습니다. 이 원고의 수치는 잘된 것만 고른 결과가 아니라, "
     "실패까지 기록하며 얻은 재현 가능한 측정값입니다.")
sub("3", "스마트폰 안에서 전부 처리합니다")
body("AI 모델 3종을 압축해 브라우저 안에 실었습니다. 압축 후에도 정확도 손실이 없는지 현장 데이터 전체로 "
     "재확인했고(손실 0), 서버가 없으니 촬영 화면이 폰 밖으로 나가지 않아 이용자 프라이버시와 관내 보안 "
     "걱정이 원천적으로 없습니다. 도서관은 웹 주소 하나로 서비스를 도입합니다.")
sub("4", "한계도 밝힙니다")
body("한 번에 읽을 수 있는 범위는 서가 1~2단(가까이 촬영)입니다. 여러 단을 멀리서 한 화면에 담으면 "
     "라벨 글자가 너무 작아져 30% 정도만 읽힙니다 — 이는 소프트웨어가 아니라 카메라 화소의 물리적 "
     "한계입니다. 그래서 촬영 가이드(한 컷에 1~2단), 먼 서가는 카메라 줌 활용, 화면 밖 책은 방향 안내로 "
     "대응하도록 설계했습니다.")

# ═══════════ Ⅴ ═══════════
heading("Ⅴ", "추진 경과")
add_table(["주차", "한 일", "결과"], [
    ["1~2주", "실증관 서가 촬영, 기준 성능 측정, 새 서가 구간 일반화 확인", "구간 추가 시 수정 0~2줄"],
    ["3~4주", "AI 검출 모델 자체 훈련, 파이프라인 통합", "규칙 방식 대비 전 지형 우세"],
    ["5주", "걸으며 스캔 검증, 판독 논리 개선, 공식 채점", "판독 정확도 99%"],
    ["6주", "모델 압축·스마트폰 이식, 공개 배포", "공개 데모 운영 시작"],
    ["7주", "현직 사서와 현장 개선: 사진 점검·점검 기록·도서찾기 확률 안내·위험도 분석", "사진 1장 16권/8초, "
     "위험도 예측 9곳 중 7 적중"],
], widths=[18, 84, 48], size=9.5)
body("모든 변경은 자동 검사 13종을 전부 통과할 때만 배포했고, 효과 없는 방법(재학습 5건 등)은 실험 "
     "기록과 함께 버렸습니다.")

# ═══════════ Ⅵ ═══════════
heading("Ⅵ", "기대효과 및 발전 방향")
sub("1", "기대효과")
bullet("이용자 — 청구기호를 몰라도 책을 찾고, 서가 앞에서 추천도서·인기대출을 발견합니다. '목록엔 있는데 "
       "서가에 없는 책'이 줄고, 어린이·어르신·도서관 초심자의 정보 접근 문턱이 낮아집니다.")
bullet("사서 — 서가 한 줄 점검이 '눈으로 수십 분'에서 '30초 걷기'로 바뀝니다. 점검이 가벼워지면 주기가 "
       "짧아지고, 이용자가 책을 못 찾는 시간도 함께 줄어듭니다.")
bullet("도서관 — 장비 투자 없이 웹 주소 하나로 도입합니다. 쌓이는 점검 기록은 구간별 오배열 통계라는, "
       "지금까지 없던 운영 데이터가 됩니다.")
sub("2", "발전 방향")
bullet("전국 확장: 정보나루 Open API로 도서관별 장서 데이터만 연결하면 즉시 적용 — 실증에서 새 서가 구간 "
       "추가에 수정 0~2줄임을 확인했습니다.")
bullet("점검 우선순위 추천: 본 실증의 위험도 분석(9곳 중 7 적중)을 정보나루 전국 대출 빅데이터와 결합해 "
       "'어느 서가부터 점검할지'를 추천하는 분석으로 키웁니다.")
bullet("실사용 파일럿: 실증관 사서와 주 1회 정기 점검 파일럿으로 점검 시간·발견 건수를 실사용 수치로 "
       "확보합니다.")
bullet("자관 시스템 연동(도입 단계): 지금은 장서 데이터 스냅샷을 쓰지만, 도서관이 정식 도입하면 자관 "
       "전산과 연동해 대출 상태까지 실시간으로 맞춰 봅니다.")
bullet("접근성 확장: 시각 약자를 위한 음성 안내('세 번째 칸, 왼쪽에서 다섯 번째') 등 폰 안에서 전부 "
       "처리되는 특성을 살린 기능.")
page_break()

# ═══════════ 참고문헌 ═══════════
heading("참고문헌", "")
doc.paragraphs[-1].runs[0].text = "참고문헌"
for ref in [
    "국립중앙도서관. (2026). 2026 도서관 데이터 활용 공모전 안내문. 국립중앙도서관.",
    "국립중앙도서관. (2026). 사서추천도서 Open API. https://nl.go.kr",
    "도서관 정보나루. (2026). 도서관 빅데이터 Open API(인기대출도서). https://www.data4library.kr",
    "영등포구립도서관. (2026). 대림도서관 자료실 안내 및 소장자료 검색. https://www.ydplib.or.kr",
    "Jocher, G., et al. (2024). Ultralytics YOLO (Version 11/26) [Computer software]. https://github.com/ultralytics/ultralytics",
    "PaddlePaddle. (2025). PaddleOCR: Awesome multilingual OCR toolkits (PP-OCRv5) [Computer software]. https://github.com/PaddlePaddle/PaddleOCR",
    "Microsoft. (2025). ONNX Runtime Web [Computer software]. https://onnxruntime.ai",
]:
    p = para(ref, size=10, space_after=4)
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.first_line_indent = Pt(-20)

# ═══════════ 부록 ═══════════
para()
heading("부록", "공개 데모 및 사용 AI 기술 명세")
sub("1", "공개 데모")
body(f"주소: {DEMO_URL} (스마트폰·PC 브라우저에서 즉시 실행, 설치 불필요). '샘플 서가' 버튼으로 실증관 "
     "서가 판독 전 과정을, '서가 점검하기'로 실시간 카메라 검출을 체험할 수 있습니다. 촬영 영상은 "
     "전송되지 않습니다.")
add_image(qr_path, 26)
sub("2", "사용 AI 기술 명세 (요강의 AI 기술 명시 의무 준수)")
add_table(["구분", "기술", "용도"], [
    ["라벨 위치 검출", "Ultralytics YOLO 경량 모델(nano)을 자체 구축 데이터 7,000여 박스로 전이학습",
     "서가 화면에서 청구기호 라벨 위치 탐지 (ONNX 변환 탑재)"],
    ["문자 인식", "PaddleOCR korean PP-OCRv5 인식 모델을 실측 라벨 702줄+합성 데이터로 파인튜닝",
     "청구기호 판독 (반정밀도 압축 — 현장 평가셋 전수 비교로 손실 0 검증)"],
    ["텍스트 줄 검출", "PaddleOCR PP-OCRv5 mobile 검출 모델(사전학습 그대로)", "라벨 안의 글줄 분리"],
    ["실행 환경", "ONNX Runtime Web (WebGPU/WASM)", "브라우저 온디바이스 추론"],
    ["판정 논리", "청구기호 퍼지 매칭 + 제목 이중 대조 + 최장증가부분수열(LIS) 순서 판정", "장서 대조·오배열 지목"],
    ["개발 보조", "AI 코딩 어시스턴트(Anthropic Claude)", "코드 작성·실험 자동화 보조 — 실험 설계·판정·실측은 "
     "본 팀이 수행·검증"],
], widths=[26, 66, 58], size=9.5)
sub("3", "산출물")
bullet("공개 웹 데모와 소스, 검출·인식 모델, 채점 스크립트, 자동 검사 13종")
bullet("실증 데이터: 전 장서 20,944권 대조 데이터, 수집 조각 5,200여 개, 사람 검증 정답 788줄(현장 평가셋), "
       "구간 위험도 분석(177구간)")
bullet("2차 발표심사 시 '분석에 활용한 데이터 및 모델 1식' 제출 가능")

doc.save(str(OUT))
print(f"저장: {OUT}")
print("다음 단계: 한글(HWP)에서 열어 분량(10p 내외)·서식 확인")
